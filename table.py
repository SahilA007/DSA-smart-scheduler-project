from docx import Document
from docx.enum.section import WD_ORIENTATION # used to set the orientation of the docx page
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT # used to center the text inside the header cells
from docx.shared import Inches, Pt, RGBColor

from docx.oxml.ns import nsdecls #needed to change the table cell bg color
from docx.oxml import parse_xml #needed to change the table cell bg color

from copy import deepcopy
from datetime import datetime


def get_template(template_name):
    template = Document(template_name)
    if template:
        return template
    else:
        raise ValueError("Invalid template name")
    

def generate_tables_template(template_name,n,departement_name):
    #Creating a new document for our tables
    output_doc = Document()
    
    #accessing the docx sections to modify page properties
    section = output_doc.sections[0]
    #setting the page orientation to landscape
    section.orientation = WD_ORIENTATION.LANDSCAPE
    #setting the page width & height
    section.page_width = Inches(13.89)
    section.page_height = Inches(12.84)

    #setting the page padding to zero
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    # setting up the font type & size for the output doc
    style = output_doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    template = get_template(template_name)
    temp_tables = template.tables
    for i in range(n):
        for table in temp_tables:
            new_table = deepcopy(table._tbl)
            paragraph = output_doc.add_paragraph()
            paragraph._p.addnext(new_table)
        # This condition is used to prevent adding an empty break page at the end of the docx file
        if i == n-1:
            break
        output_doc.add_page_break()
    doc_name = f"Generated_{departement_name}_timetables_"+ str(datetime.now().year-1)+"_"+str(datetime.now().year)+".docx"
    output_doc.save(doc_name)
    return doc_name


def fill_tables(matrix,doc,departement_name,semester_name,start_date,end_date):
    headers=[]
    timetables =[]
    template = Document(doc) 
    tables = template.tables
    i=0 #header 
    j=1 #timetable
    group = 0
    while i<j and j< len(tables):
        
        fill_header(tables[i],group+1,departement_name,semester_name,start_date,end_date)
        headers.append(tables[i])
        fill_timetable(matrix,tables[j],group)
        timetables.append(tables[j])
        
        i+=2
        j+=2
        group+=1

    template.save(doc)

def fill_header(table,group,departement_name,semester_name,start_date,end_date):
    append_text_to_header(table.cell(1,0),f"{departement_name}")
    append_text_to_header(table.cell(1,1),f"Class {group}",True)
    append_text_to_header(table.cell(0,3),f"{datetime.today().year-1} / {datetime.today().year}")
    append_text_to_header(table.cell(1,3),f"{semester_name}")
    append_text_to_header(table.cell(2,2),f"{datetime.today().strftime('%Y/%m/%d %I:%M:%S %p')}")
    append_text_to_header(table.cell(3,3),f"{start_date}/{end_date}")
    
def append_text_to_header(cell,text,bold=False):
    run = cell.paragraphs[0].add_run(text)
    run.font.bold= bold
    run.alignment = WD_PARAGRAPH_ALIGNMENT

def fill_timetable(matrix,table,group):           
    for r,row in enumerate(table.rows[1:]):    
        update= False
        for col,cell in enumerate(row.cells[1:]):
            if cell.text.strip() == "":
                if matrix[group][r][col] !="*":
                    append_text_to_cell(cell,matrix[group][r][col].course+"\n",11,(16, 14, 22))
                    append_text_to_cell(cell,matrix[group][r][col].prof+"\n",10,(16, 14, 22))
                    append_text_to_cell(cell,matrix[group][r][col].classroom,9,(16, 14, 22))
                    set_cell_bg_color(cell,matrix[group][r][col].color)
                    col+=1
                    update = True
                else:
                    continue
                
        if update:
            r+=1

def append_text_to_cell(cell,text,font_size,color):
    run = cell.paragraphs[0].add_run(text)
    run.font.bold = True
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(*color)


# set the table cell bg color based on the course type
def set_cell_bg_color(cell,bgcolor):
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), bgcolor))
    cell._element.tcPr.append(shading_elm)


def show_header_tables(headers):
    for _ in range(len(headers)):
        for header in headers:
            for rnum,row in enumerate(header.rows):
                print(f"row:{rnum}")
                for num,cell in enumerate(row.cells):
                    print(f"cell:{num} {cell.text}",end =" | ")
                print("\n------------------")
            print()
            break
        break

#the last two function will not make it to the github version:

def show_timetables(timetables):
    for _ in range(len(timetables)):
        for timetable in timetables:
            for rnum,row in enumerate(timetable.rows):
                print(f"row:{rnum}")
                for num,cell in enumerate(row.cells):
                    print(f"cell:{num} {cell.text}",end =" | ")
                print("\n------------------")
            print()
            
    